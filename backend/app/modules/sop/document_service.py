from __future__ import annotations

import io
from uuid import UUID

import docx
from pypdf import PdfReader
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.core.exceptions import LLMError
from app.modules.sop.models import Document, DocumentChunk
from app.infrastructure.llm.openrouter import get_openrouter_provider
from app.modules.sop.repository import DocumentRepository
from app.modules.sop.schemas import DocumentRead

_CHUNK_SIZE = 1000
_CHUNK_OVERLAP = 150
_ALLOWED_EXTENSIONS = {".pdf", ".txt", ".docx"}


class DocumentValidationError(Exception):
    """Client-input error (bad file size/type) -- routes map this to 400,
    not the global EduPathError->503 handler used for external-service
    failures."""


def _extract_text(filename: str, raw: bytes) -> str:
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if suffix == "pdf":
        reader = PdfReader(io.BytesIO(raw))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == "docx":
        document = docx.Document(io.BytesIO(raw))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
    if suffix == "txt":
        return raw.decode("utf-8", errors="replace")
    raise DocumentValidationError(f"Unsupported file type: .{suffix or 'unknown'}. Allowed: {sorted(_ALLOWED_EXTENSIONS)}")


def _chunk_text(text: str, *, size: int = _CHUNK_SIZE, overlap: int = _CHUNK_OVERLAP) -> list[str]:
    text = text.strip()
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + size
        chunks.append(text[start:end].strip())
        start = end - overlap
    return [chunk for chunk in chunks if chunk]


class DocumentService:
    def __init__(self, repository: DocumentRepository | None = None, provider=None) -> None:
        self._repository = repository or DocumentRepository()
        self._provider = provider or get_openrouter_provider()

    async def upload(
        self, session: AsyncSession, *, profile_id: UUID, filename: str, document_type: str, raw: bytes
    ) -> DocumentRead:
        max_size_bytes = int(settings.max_document_size_mb * 1024 * 1024)
        if len(raw) > max_size_bytes:
            raise DocumentValidationError(f"File exceeds the {settings.max_document_size_mb}MB limit.")

        text = _extract_text(filename, raw)
        document = Document(profile_id=profile_id, filename=filename, document_type=document_type, content_text=text)
        document = await self._repository.create(session, document)

        for index, chunk_text in enumerate(_chunk_text(text)):
            embedding = None
            try:
                embedding = self._provider.embed_text(chunk_text, model=settings.embedding_model)
            except LLMError:
                # A document is still useful (e.g. for direct SOP grounding
                # via its raw content_text) even if embeddings are
                # temporarily unavailable -- don't fail the whole upload.
                pass
            session.add(DocumentChunk(document_id=document.id, chunk_index=index, content=chunk_text, embedding=embedding))

        await session.commit()
        await session.refresh(document)
        document = await self._repository.get(session, document.id)
        return self._to_read(document)

    async def list_for_profile(self, session: AsyncSession, profile_id: UUID) -> list[DocumentRead]:
        documents = await self._repository.list_for_profile(session, profile_id)
        return [self._to_read(document) for document in documents]

    async def delete(self, session: AsyncSession, document_id: UUID) -> bool:
        document = await self._repository.get(session, document_id)
        if document is None:
            return False
        await self._repository.delete(session, document)
        return True

    async def retrieve_relevant_chunks(
        self, session: AsyncSession, profile_id: UUID, query_text: str, *, limit: int = 5
    ) -> list[str]:
        """Used by grounded_context()/SOP generation to ground on the
        student's actual uploaded documents instead of just profile fields."""
        try:
            embedding = self._provider.embed_text(query_text, model=settings.embedding_model)
        except LLMError:
            return []
        chunks = await self._repository.search_similar_chunks(session, embedding, profile_id=profile_id, limit=limit)
        return [chunk.content for chunk in chunks]

    @staticmethod
    def _to_read(document: Document) -> DocumentRead:
        return DocumentRead(
            id=str(document.id),
            profile_id=str(document.profile_id) if document.profile_id else None,
            filename=document.filename,
            document_type=document.document_type,
            chunk_count=len(document.chunks),
            created_at=document.created_at,
        )
